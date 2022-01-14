import collections
import re
from typing import NamedTuple
import logging

import docx

from ifaxbotcovid.bot.helpers import FileSaver


class Sender:
    '''
    A wrapper around bot.sendmessage, provides
    useful methods to send ready answer, logs
    and/or warningmessage according to user request,
    and answer cooked by CovidChef.
    '''

    def __init__(self, bot, message, answer, logger,
                 logrequest: bool = False):
        self.bot = bot
        self.message = message
        self.answer = answer
        self.botlogger = logger
        self.logrequest = logrequest

    def send_warn(self):
        if self.answer.warnmessage:
            self.bot.send_message(
                self.message.chat.id, self.answer.warnmessage)
            self.botlogger.info(
                'Warning message sent to %s' % self.message.from_user.username
            )

    def send_log(self):
        if self.logrequest:
            self.bot.send_message(self.message.chat.id, self.answer.log)
            self.botlogger.info(
                'Log message sent to %s' % self.message.from_user.username)

    def send_directly(self):
        self.send_warn()
        self.bot.send_message(self.message.chat.id, self.answer.ready_text)
        self.botlogger.info(
            'Ready answer sent to %s' % self.message.from_user.username
        )
        self.send_log()

    def send_asfile(self):
        path = FileSaver.to_file(
            text=self.answer.ready_text,
            username=self.message.from_user.username
        )
        self.send_warn()
        self.send_log()
        try:
            self.bot.send_document(
                self.message.chat.id, open(path), 'document'
            )
        except Exception as exc:
            self.botlogger.error(
                'No system log file found! Exception: %s' % exc)
            self.bot.send_message(
                self.message.chat.id, 'Ошибка при отправке файла')


class CommandParser:

    logger = logging.getLogger(__name__)

    class CommandList(NamedTuple):
        asfile: bool = False   # if user requested answer in a file
        logrequest: bool = False
        short: int = 0     # var responsible to shorten region list

    def _get_short(text):
        short = 0
        short_pattern = re.compile(r'\$\$ {0,1}(\d{2,4})')
        match = short_pattern.search(text)
        if match:
            result_as_text = match.group(1)
            CommandParser.logger.debug(result_as_text)
            try:
                short = int(result_as_text)
            except ValueError:
                CommandParser.logger.warning(' '.join((
                    'CommandParser failed to get integer',
                    'from a match object. $short set to 0'
                    ))
                )
        return short

    def get_settings(message) -> CommandList:
        text = message.text.lower()
        logreq_condition = 'йй' in text
        asfile_condition = '$$' in text
        short = CommandParser._get_short(text)
        return CommandParser.CommandList(asfile=asfile_condition,
                                         logrequest=logreq_condition,
                                         short=short)


class DocxReader:

    Tables = collections.namedtuple(
        'Tables', ['new_cases', 'new_deaths', 'recovered']
    )

    def __init__(self, file):
        self.doc = docx.Document(file)

    def to_text(self):
        result = self.read()
        return result

    def read(self):
        assert len(self.doc.tables) == 3, (
            f'Expected 3 tables, got {len(self.doc.tables)}'
        )
        tbls = [self.read_table(table) for table in self.doc.tables]
        tables = DocxReader.Tables(*tbls)
        pargs = [parg.text for parg in self.doc.paragraphs]
        text = self.construct(pargs, tables)
        return text

    @staticmethod
    def construct(pargs, tables):
        mapping = {
            'Распределение по субъектам': tables.new_cases,
            'подтверждено': tables.new_deaths,
            'выписано по выздоровл': tables.recovered
        }
        for i, parg in enumerate(pargs):
            for key in mapping:
                if key in parg:
                    pargs.insert(i + 1, mapping[key])

        return '\n'.join(pargs)

    @staticmethod
    def read_table(table, delim='\t', newline='\n'):
        if not table:
            return
        rows = [
            delim.join(
                [cell.text for cell in row.cells]
            ) for row in table.rows
        ]
        return newline.join(rows)
