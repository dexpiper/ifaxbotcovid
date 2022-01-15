from collections import deque, namedtuple
import re
from time import time
import os
import logging
from typing import NamedTuple
from pathlib import Path

import docx


class MessageStorage():
    '''
    Stores sequentional messages sent by users and chat_id's,
        concatenates sequentional messages from single user if
        they were sent rapidly one after another (as when Telegram
        cuts apart one long message into pieces and sends them
        together).

    Usage:
        storage = MessageStorage()
        storage.append(text='Some message', chat_id='12345')
        joint_message = storage.get_joint()
        if joint_message.valid:
            text = joint_message.text

    Parameters:

        maxlen - maximum messages to form a joint message (default 3);
        check_phrases - (optional) phrases to be in joint message to mark
            it valid;
        stop_phrase - (optional) phrase that must be in the last message. If
            defined, no output would be given and bo messages would be glued
            unless that phrase is in the last message.
        time_gap - (optional) two or more sequentional messages from one user
            whould be considered parts of one single message if time between
            them is less or equal that parameter. Default - 1.5 (seconds)
    '''
    class InnerMessage(NamedTuple):
        time: int   # time when message came
        text: str     # message text
        chat_id: int  # chat_id of the user

    def __init__(self, maxlen: int = 3, check_phrases: list = [],
                 stop_phrase: str = '', time_gap: float = 1.5):
        self.maxlen = maxlen
        self._spawn()
        self.check_phrases = check_phrases
        self.stop_phrase = stop_phrase
        self.time_gap = time_gap

    def _spawn(self):
        self._db = deque(maxlen=self.maxlen*2)
        self.append()

    def __len__(self):
        return len(self._db)

    def append(self, text='', chat_id='', time=int(time())) -> None:
        self._db.append(
            MessageStorage.InnerMessage(
                time=time,
                text=text,
                chat_id=chat_id
            )
        )

    def drop(self):
        self._spawn()

    def _last(self) -> InnerMessage:
        return self._db[-1]

    def _last_before_last(self) -> InnerMessage:
        try:
            return self._db[-2]
        except IndexError:
            return ''

    def _get_sequence(self):
        '''
        Return all messages sent by the user
        who wrote the last message.
        '''
        last_message = self._last()
        return tuple(message for message in self._db if
                     message.chat_id == last_message.chat_id)

    def get_joint(self):
        '''
        Concatenate sequenctional messages from
        user who sent the last message if time gap
        between his messages is less then given time gap.

        Returns JointMessage object
        '''
        last = self._last()
        if (self.stop_phrase) and (
            self.stop_phrase not in last.text
        ):
            return JointMessage('', self, invalid_flag=True)

        seq = self._get_sequence()          # messages from user in storage
        texts = []                          # list of messages to concatenate
        t = last.time

        for i in range(-1, -(len(seq) + 1), -1):
            # iterating -1 step backwards, from the last
            # message in sequence to the first
            msg = seq[i]
            time_difference = t - msg.time  # time between messages
            if time_difference <= self.time_gap:
                # inserting message from the beginning
                texts.insert(0, msg.text)
            t = msg.time

        glued_text = ''.join(texts)
        return JointMessage(glued_text, self)


class JointMessage():

    def __init__(self,
                 message: str,
                 message_storage: MessageStorage,
                 invalid_flag: bool = False):
        self.text = ''
        self._message = message
        self._flag = invalid_flag
        self.check_phrases = message_storage.check_phrases
        self.stop_phrase = message_storage.stop_phrase
        if invalid_flag:
            self.valid = False
        else:
            self.valid = self.validate()
        if self.valid:
            self.text = message

    def __eq__(self, other):
        if self._message == other:
            return True
        else:
            return False

    def __len__(self):
        return len(self._message)

    def __repr__(self):
        if self.valid:
            v = 'Valid'
        else:
            v = 'Not valid'
        return f"<JointMessage ({v}), length: {self.__len__()}>"

    def validate(self):
        if self.check_phrases:
            for phrase in self.check_phrases:
                if phrase not in self._message:
                    return False
        if self.stop_phrase:
            if self.stop_phrase not in self._message:
                return False
        return True


class LogConstructor:

    def join_log_message(log: list):
        '''
        Construct a single string to send to the user
        as a log message (parsing results)
        '''
        result = ''
        try:
            for el in log:
                if type(el) == list:
                    for item in el:
                        result += (item + '\n')
                else:
                    result += (el + '\n')
        except Exception as exc:
            return f'Exception raised during log construction: {exc}'
        return result


class FileSaver:

    def to_file(text: str,
                username: str = 'default',
                timestamp: int = int(time()),
                tempdir: str = './temp'
                ) -> str:
        path = f'{tempdir}/{username}{str(timestamp)}.txt'
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, 'w') as f:
            f.write(text)
        return path

    def from_file(contents: str,
                  filename: str,
                  username: str = 'default',
                  timestamp: int = int(time()),
                  tempdir: str = './temp'
                  ) -> str:
        path = f'{tempdir}/{username}{str(timestamp)}.{filename[-4:]}'
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, 'wb') as f:
            f.write(contents)
        return path

    def del_file(path: str):
        path = Path(path)
        if path.is_file():
            os.remove(path)


class Sender:
    '''
    A wrapper around bot.sendmessage, provides
    useful methods to send ready answer, logs
    and/or warningmessage according to user request,
    and answer cooked by CovidChef.
    '''

    def __init__(self, bot, answer: object,
                 logger=logging.getLogger(__name__),
                 logrequest: bool = False):
        self.bot = bot
        self.answer = answer
        self.message = answer.message_object
        self.botlogger = logger
        self.logrequest = logrequest
        self.path = None

    def send_warn(self):
        if self.answer.warnmessage:
            self.bot.send_message(
                self.message.chat.id, '\u2757' + self.answer.warnmessage)
            self.botlogger.info(
                'Warning message sent to %s' % self.message.from_user.username
            )

    def send_log(self):
        if self.logrequest:
            self.bot.send_message(self.message.chat.id, self.answer.log)
            self.botlogger.info(
                'Log message sent to %s' % self.message.from_user.username)

    def send_directly(self):
        self.send_warn()
        self.bot.send_message(self.message.chat.id, self.answer.ready_text)
        self.botlogger.info(
            'Ready answer sent to %s' % self.message.from_user.username
        )
        self.send_log()

    def send_asfile(self):
        self.path = FileSaver.to_file(
            text=self.answer.ready_text,
            username=self.message.from_user.username
        )
        self.send_warn()
        self.send_log()
        try:
            self.bot.send_document(
                self.message.chat.id, open(self.path), 'document'
            )
        except Exception as exc:
            self.botlogger.error(
                'File not found! Exception: %s' % exc)
            self.bot.send_message(
                self.message.chat.id, 'Ошибка при отправке файла')


class CommandParser:

    logger = logging.getLogger(__name__)

    class CommandList(NamedTuple):
        asfile: bool = False   # if user requested answer in a file
        logrequest: bool = False
        short: int = 0     # var responsible to shorten region list

    def _get_short(text):
        short = 0
        short_pattern = re.compile(r'\$\$ {0,1}(\d{2,4})')
        match = short_pattern.search(text)
        if match:
            result_as_text = match.group(1)
            CommandParser.logger.debug(result_as_text)
            try:
                short = int(result_as_text)
            except ValueError:
                CommandParser.logger.warning(' '.join((
                    'CommandParser failed to get integer',
                    'from a match object. $short set to 0'
                    ))
                )
        return short

    def get_settings(message) -> CommandList:
        text = message.text.lower()
        logreq_condition = 'йй' in text
        asfile_condition = '$$' in text
        short = CommandParser._get_short(text)
        return CommandParser.CommandList(asfile=asfile_condition,
                                         logrequest=logreq_condition,
                                         short=short)


class DocxReader:

    Tables = namedtuple(
        'Tables', ['new_cases', 'new_deaths', 'recovered']
    )

    def __init__(self, file):
        self.doc = docx.Document(file)

    def to_text(self):
        result = self.read()
        return result

    def read(self):
        pargs = [parg.text for parg in self.doc.paragraphs]
        if not self.doc.tables:
            text = self.construct(pargs)
            return text
        assert len(self.doc.tables) == 3, (
            f'Expected 3 tables, got {len(self.doc.tables)}'
        )
        tbls = [self.read_table(table) for table in self.doc.tables]
        tables = DocxReader.Tables(*tbls)
        text = self.construct(pargs, tables)
        return text

    @staticmethod
    def construct(pargs, tables=None):
        """
        Joining paragraph objects as plain text.
        Inserting text tables in correspondent places if there is any.
        """
        checkphrases = ['Распределение по субъектам', 'подтверждено',
                        'выписано по выздоровл', 'COVID-19', 'смерт',
                        'случа', 'зарегистрир']
        if not tables:
            text = '\n'.join(pargs)
            assert all(
                [True if key in text else False for key in checkphrases]
            )
            return text

        mapping = {
            'Распределение по субъектам': tables.new_cases,
            'подтверждено': tables.new_deaths,
            'выписано по выздоровл': tables.recovered
        }
        for i, parg in enumerate(pargs):
            for key in mapping:
                if key in parg:
                    pargs.insert(i + 1, mapping[key])

        text = '\n'.join(pargs)
        assert all(
            [True if key in text else False for key in checkphrases]
        )
        return text

    @staticmethod
    def read_table(table, delim='\t', newline='\n'):
        if not table:
            return
        rows = [
            delim.join(
                [cell.text for cell in row.cells]
            ) for row in table.rows
        ]
        return newline.join(rows)
